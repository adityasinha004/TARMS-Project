from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os

def create_complete_project_file():
    doc = Document()
    
    # Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set default alignment to Justify
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for i in range(1, 5):
        h_style = doc.styles[f'Heading {i}']
        h_font = h_style.font
        h_font.name = 'Arial'
        h_font.bold = True
        h_font.color.rgb = RGBColor(0, 0, 0)
        h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT # Headings left aligned
        if i == 1: h_font.size = Pt(16)
        elif i == 2: h_font.size = Pt(14)
        elif i == 3: h_font.size = Pt(12)
        elif i == 4: # Special Blue Style for Use Cases
            h_font.size = Pt(12)
            h_font.italic = True
            h_font.color.rgb = RGBColor(79, 129, 189)

    # --- TITLE PAGE ---
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('UCS503 - Software Engineering Lab')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(16)
    p.runs[0].bold = True
    
    p = doc.add_paragraph('TARMS')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(24)
    p.runs[0].bold = True
    
    p = doc.add_paragraph('UCS 503 Software Engineering Project Report')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('END-Semester Evaluation')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].bold = True
    
    doc.add_paragraph('\n\n')
    
    p = doc.add_paragraph('Submitted by:')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    p = doc.add_paragraph('[Student Name] - [Roll No]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('Group No: [Group ID]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    
    p = doc.add_paragraph('Submitted to: Dr ASHIMA SINGH')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('Computer Science and Engineering Department\nTIET, Patiala\nApril 2025')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # --- TABLE OF CONTENTS ---
    doc.add_heading('TABLE OF CONTENTS', level=1)
    toc = [
        "1. Project Selection Phase",
        "   1.1 Software Bid",
        "   1.2 Project Overview",
        "2. Analysis Phase",
        "   2.1 Use Cases",
        "       2.1.1 Use-Case Diagrams",
        "       2.1.2 Use Case Templates",
        "   2.2 Activity Diagram and Swimlane Diagrams",
        "   2.3 Data Flow Diagrams (DFDs)",
        "       2.3.1 DFD Level 0",
        "       2.3.2 DFD Level 1",
        "       2.3.3 DFD Level 2",
        "   2.4 Software Requirement Specification in IEEE Format",
        "   2.5 User Stories and Story Cards",
        "3. Design Phase",
        "   3.1 Class Diagram",
        "   3.2 Sequence Diagram",
        "   3.3 Collaboration Diagram",
        "   3.4 State Chart Diagrams",
        "4. Implementation",
        "   4.1 Component Diagrams",
        "   4.2 Deployment Diagrams",
        "   4.3 Screenshots",
        "5. Testing",
        "   5.1 Test Plan",
        "   5.2 Test Cases",
        "   5.3 Test Reports"
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT # TOC Left Aligned
    doc.add_page_break()

    # --- 1. PROJECT SELECTION PHASE ---
    doc.add_heading('1. Project Selection Phase', level=1)
    
    doc.add_heading('1.1 Software Bid', level=2)
    p = doc.add_paragraph()
    p.add_run('Software Bid/ Project Teams\nUCS 503- Software Engineering Lab').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Group: [Group ID]\t\t\tDated: [Date]')
    doc.add_paragraph('Team Name: [Team Name]')
    doc.add_paragraph('Team ID: [ID]')
    
    # Members Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Name'
    hdr[1].text = 'Roll No'
    hdr[2].text = 'Project Experience'
    hdr[3].text = 'Programming Language'
    
    # Placeholder row
    row = table.add_row().cells
    row[0].text = '[Student Name]'
    row[1].text = '[Roll No]'
    row[2].text = 'Web Dev'
    row[3].text = 'Python/JS'
    
    doc.add_paragraph('\nProgramming Language Preference: 1. Web Dev (Node.js) 2. Python 3. SQL')
    
    doc.add_paragraph('\nChoices of Projects:')
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Project Name'
    table.rows[0].cells[1].text = 'Unique Selling Point'
    table.rows[1].cells[0].text = '1. TARMS'
    table.rows[1].cells[1].text = 'Streamlines auditorium booking, reducing manual errors and conflicts.'
    
    doc.add_heading('1.2 Project Overview', level=2)
    doc.add_paragraph("The goal is to design a software for Thapar students/faculty to manage auditorium bookings. TARMS (Thapar Auditorium and Room Management System) allows users to view availability, book rooms, and track status. Admins/Faculty can approve requests.")
    doc.add_page_break()

    # --- 2. ANALYSIS PHASE ---
    doc.add_heading('2. Analysis Phase', level=1)
    
    doc.add_heading('2.1 Use Cases', level=2)
    doc.add_heading('2.1.1 Use-Case Diagrams', level=3)
    if os.path.exists('use_case_diagram.png'):
        doc.add_picture('use_case_diagram.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif os.path.exists('image1.png'):
        doc.add_picture('image1.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Use Case Diagram]")

    doc.add_heading('2.1.2 Use Case Templates', level=3)
    
    def add_uc(title, actor, desc, pre, flow, post):
        doc.add_heading(f'Use Case: {title}', level=4)
        doc.add_paragraph(f'Actor: {actor}')
        doc.add_paragraph(f'Description: {desc}')
        doc.add_paragraph(f'Pre-conditions: {pre}')
        doc.add_paragraph(f'Main Flow: {flow}')
        doc.add_paragraph(f'Post-conditions: {post}')
        doc.add_paragraph()

    add_uc('Student Login', 'Student', 'Access system', 'Registered', '1. Enter creds 2. Validate 3. Dashboard', 'Logged in')
    add_uc('Book Room', 'Student', 'Request room', 'Logged in', '1. Select Room 2. Fill Form 3. Submit', 'Request Pending')
    add_uc('Approve Request', 'Faculty', 'Approve booking', 'Logged in', '1. View Request 2. Approve/Reject', 'Status Updated')

    doc.add_heading('2.2 Activity Diagram and Swimlane Diagrams', level=2)
    doc.add_heading('2.2.1 Activity Diagram', level=3)
    if os.path.exists('activity_booking.png'):
        doc.add_picture('activity_booking.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('2.2.2 Swimlane Diagram', level=3)
    if os.path.exists('swimlane_diagram.png'):
        doc.add_picture('swimlane_diagram.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('2.3 Data Flow Diagrams (DFDs)', level=2)
    doc.add_heading('2.3.1 DFD Level 0', level=3)
    if os.path.exists('dfd_level_0.png'):
        doc.add_picture('dfd_level_0.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_heading('2.3.2 DFD Level 1', level=3)
    if os.path.exists('dfd_level_1.png'):
        doc.add_picture('dfd_level_1.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_heading('2.3.3 DFD Level 2', level=3)
    if os.path.exists('dfd_level_2.png'):
        doc.add_picture('dfd_level_2.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('2.4 Software Requirement Specification in IEEE Format', level=2)
    doc.add_paragraph("1. Introduction\n1.1 Purpose: Define requirements for TARMS.\n1.2 Scope: Web-based room booking.\n2. Overall Description\n2.1 Product Perspective: Replaces manual system.\n2.2 User Characteristics: Students, Faculty.\n3. Specific Requirements\n3.1 Functional: Authentication, Booking, Approval.\n3.2 Non-Functional: Performance, Security.")

    doc.add_heading('2.5 User Stories and Story Cards', level=2)
    
    def add_story(id, title, story, success, failure):
        doc.add_paragraph(f'{id} {title}').runs[0].bold = True
        doc.add_paragraph(story)
        doc.add_paragraph('CONFIRMATION').runs[0].bold = True
        doc.add_paragraph(f'1. Success: {success}')
        doc.add_paragraph(f'2. Failure: {failure}')
        doc.add_paragraph()

    add_story('#0001', 'STUDENT LOGIN', 'As a registered user, I want to log in, so I can book rooms.', 'Valid user logged in.', 'Displays "Invalid credentials".')
    add_story('#0002', 'BOOK ROOM', 'As a student, I want to book a room, so I can host an event.', 'Request submitted.', 'Error "Room not available".')
    add_story('#0003', 'APPROVE REQUEST', 'As faculty, I want to approve requests, so rooms are allocated.', 'Status updated to Approved.', 'Error updating status.')

    doc.add_page_break()

    # --- 3. DESIGN PHASE ---
    doc.add_heading('3. Design Phase', level=1)
    
    doc.add_heading('3.1 Class Diagram', level=2)
    if os.path.exists('class_diagram.png'):
        doc.add_picture('class_diagram.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.2 Sequence Diagram', level=2)
    doc.add_heading('3.2.1 Sequence diagram to Book Room', level=3)
    if os.path.exists('sequence_add_booking.png'):
        doc.add_picture('sequence_add_booking.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif os.path.exists('sequence_booking.png'):
        doc.add_picture('sequence_booking.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.2.2 Sequence diagram to View Room', level=3)
    if os.path.exists('sequence_view_room.png'):
        doc.add_picture('sequence_view_room.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.2.3 Sequence diagram for Login', level=3)
    if os.path.exists('sequence_login.png'):
        doc.add_picture('sequence_login.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.3 Collaboration Diagram', level=2)
    doc.add_heading('3.3.1 Collaboration diagram to Book Room', level=3)
    if os.path.exists('collab_add_booking.png'):
        doc.add_picture('collab_add_booking.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif os.path.exists('collaboration_diagram.png'):
        doc.add_picture('collaboration_diagram.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.3.2 Collaboration Diagram to View Room', level=3)
    if os.path.exists('collab_view_room.png'):
        doc.add_picture('collab_view_room.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.4 State Chart Diagrams', level=2)
    doc.add_heading('3.4.1 State Chart Diagram for Unregistered User', level=3)
    if os.path.exists('state_unregistered.png'):
        doc.add_picture('state_unregistered.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif os.path.exists('state_chart_diagram.png'):
        doc.add_picture('state_chart_diagram.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.4.2 State Chart Diagram for Booking Request', level=3)
    if os.path.exists('state_booking.png'):
        doc.add_picture('state_booking.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- 4. IMPLEMENTATION ---
    doc.add_heading('4. Implementation', level=1)
    
    doc.add_heading('4.1 Component Diagrams', level=2)
    if os.path.exists('component_diagram.png'):
        doc.add_picture('component_diagram.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('4.2 Deployment Diagrams', level=2)
    if os.path.exists('deployment_diagram.png'):
        doc.add_picture('deployment_diagram.png', width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('4.3 Screenshots', level=2)
    doc.add_paragraph("[Login Page Screenshot]")
    doc.add_paragraph("[Dashboard Screenshot]")
    doc.add_page_break()

    # --- 5. TESTING ---
    doc.add_heading('5. Testing', level=1)
    
    doc.add_heading('5.1 Test Plan', level=2)
    doc.add_paragraph("Scope: Unit, Integration, System Testing.")
    
    doc.add_heading('5.2 Test Cases', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'ID'; hdr[1].text = 'Desc'; hdr[2].text = 'Expected'; hdr[3].text = 'Status'
    row = table.add_row().cells
    row[0].text = 'TC1'; row[1].text = 'Login'; row[2].text = 'Success'; row[3].text = 'Pass'
    
    doc.add_heading('5.3 Test Reports', level=2)
    doc.add_paragraph("All tests passed.")

    output_path = '/Users/adityasinha/Downloads/TARMS-main/TARMS_Project_File.docx'
    doc.save(output_path)
    print(f"Successfully created project file at: {output_path}")

if __name__ == "__main__":
    create_complete_project_file()
