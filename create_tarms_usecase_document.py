from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def create_tarms_usecase_document():
    # Create a new Document
    doc = Document()
    
    # Set document margins to minimize blank spaces
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Title
    title = doc.add_heading('TARMS - Use Case Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    intro_text = '''This document provides comprehensive use case documentation for the Thapar Auditorium and Room Management System (TARMS). The use cases describe the interactions between users and the system, detailing the functional requirements and business processes.'''
    doc.add_paragraph(intro_text)
    
    # Actors
    doc.add_heading('Actors', level=1)
    doc.add_heading('Primary Actors', level=2)
    primary_actors = '''• Student: Undergraduate and graduate students who request room bookings
• Faculty: Teaching staff who approve/reject booking requests
• System Administrator: IT staff who manage the system'''
    doc.add_paragraph(primary_actors)
    
    doc.add_heading('Secondary Actors', level=2)
    secondary_actors = '''• Department Coordinator: Staff who oversee room availability
• Information Technology Support: Staff who maintain system infrastructure'''
    doc.add_paragraph(secondary_actors)
    
    # Student Use Cases
    doc.add_heading('Student Use Cases', level=1)
    
    # UC-001: Student Registration
    doc.add_heading('UC-001: Student Registration', level=2)
    uc001_content = '''Actor: Student
Precondition: Student has valid university credentials
Main Flow:
1. Student navigates to the registration page
2. Student enters personal details (Name, UserID, Email, Mobile Number, DOB)
3. Student sets a password
4. Student selects member type
5. System validates the information
6. System checks for duplicate UserID and Email
7. System creates new student account
8. System displays success message
9. Student is redirected to login page

Alternative Flows:
3a. If password is weak, system prompts for stronger password
6a. If UserID/Email already exists, system displays error message
6b. If validation fails, system displays error message

Postcondition: Student account is created and student can login'''
    doc.add_paragraph(uc001_content)
    
    # UC-002: Student Login
    doc.add_heading('UC-002: Student Login', level=2)
    uc002_content = '''Actor: Student
Precondition: Student has a registered account
Main Flow:
1. Student navigates to login page
2. Student enters UserID and password
3. System validates credentials
4. System creates user session
5. System redirects to student dashboard
6. Student can access all student features

Alternative Flows:
3a. If credentials are invalid, system displays error message
3b. If account is locked, system displays appropriate message

Postcondition: Student is logged in and can access the system'''
    doc.add_paragraph(uc002_content)
    
    # UC-003: View Available Rooms
    doc.add_heading('UC-003: View Available Rooms', level=2)
    uc003_content = '''Actor: Student
Precondition: Student is logged in
Main Flow:
1. Student accesses the home dashboard
2. System retrieves all available rooms from database
3. System displays room list with details (Room Number, Room Name, Status)
4. Student can see which rooms are available for booking
5. Student can click on a room to view more details

Alternative Flows:
2a. If no rooms are available, system displays appropriate message
2b. If database error occurs, system displays error message

Postcondition: Student can see available rooms and proceed with booking'''
    doc.add_paragraph(uc003_content)
    
    # UC-004: Submit Room Booking Request
    doc.add_heading('UC-004: Submit Room Booking Request', level=2)
    uc004_content = '''Actor: Student
Precondition: Student is logged in and room is available
Main Flow:
1. Student selects a room for booking
2. Student clicks on "Request Form" button
3. System displays booking form
4. Student fills out event details:
   - Event name and purpose
   - Event date and time
   - Expected number of people
   - AC requirement
   - Sound system requirement
   - Society/club information
5. Student submits the form
6. System validates the form data
7. System creates booking request with "Requested" status
8. System sends notification to faculty
9. System displays confirmation message
10. Student can track request status

Alternative Flows:
6a. If required fields are missing, system displays validation error
6b. If date/time conflicts exist, system displays conflict message
7a. If system error occurs, booking request fails

Postcondition: Booking request is submitted and pending faculty approval'''
    doc.add_paragraph(uc004_content)
    
    # UC-005: Track Booking Status
    doc.add_heading('UC-005: Track Booking Status', level=2)
    uc005_content = '''Actor: Student
Precondition: Student has submitted booking requests
Main Flow:
1. Student navigates to "My Requests" page
2. System retrieves all booking requests for the student
3. System displays request history with status:
   - Requested: Pending faculty approval
   - Approved: Request approved by faculty
   - Rejected: Request rejected by faculty
4. Student can view detailed information for each request
5. Student can delete pending requests if needed

Alternative Flows:
2a. If no requests found, system displays appropriate message
2b. If database error occurs, system displays error message

Postcondition: Student can monitor all their booking requests'''
    doc.add_paragraph(uc005_content)
    
    # UC-006: Update Personal Information
    doc.add_heading('UC-006: Update Personal Information', level=2)
    uc006_content = '''Actor: Student
Precondition: Student is logged in
Main Flow:
1. Student navigates to "Change Information" page
2. System retrieves current student information
3. System displays editable form with current data
4. Student modifies required fields
5. Student submits updated information
6. System validates the updated data
7. System updates student profile in database
8. System displays success message

Alternative Flows:
6a. If validation fails, system displays error message
7a. If update fails, system displays error message

Postcondition: Student profile is updated with new information'''
    doc.add_paragraph(uc006_content)
    
    # UC-007: Change Password
    doc.add_heading('UC-007: Change Password', level=2)
    uc007_content = '''Actor: Student
Precondition: Student is logged in
Main Flow:
1. Student navigates to "Change Password" page
2. Student enters current password
3. Student enters new password
4. Student confirms new password
5. System validates current password
6. System validates new password strength
7. System updates password in database
8. System displays success message
9. Student is logged out and must login with new password

Alternative Flows:
5a. If current password is incorrect, system displays error
6a. If new password is weak, system prompts for stronger password
6b. If passwords don't match, system displays error

Postcondition: Student password is updated and session is terminated'''
    doc.add_paragraph(uc007_content)
    
    # UC-008: Submit Feedback
    doc.add_heading('UC-008: Submit Feedback', level=2)
    uc008_content = '''Actor: Student
Precondition: Student is logged in
Main Flow:
1. Student navigates to "Feedback" page
2. Student fills out feedback form with:
   - Rating
   - Comments
   - Suggestions
3. Student submits feedback
4. System validates feedback data
5. System stores feedback in database
6. System displays thank you message

Alternative Flows:
4a. If required fields are missing, system displays validation error
5a. If storage fails, system displays error message

Postcondition: Feedback is submitted and stored for review'''
    doc.add_paragraph(uc008_content)
    
    # Faculty Use Cases
    doc.add_heading('Faculty Use Cases', level=1)
    
    # UC-009: Faculty Registration
    doc.add_heading('UC-009: Faculty Registration', level=2)
    uc009_content = '''Actor: Faculty
Precondition: Faculty has valid university credentials
Main Flow:
1. Faculty navigates to faculty registration page
2. Faculty enters details (Name, Position, Username, Email, Mobile)
3. Faculty sets a password
4. Faculty selects member type
5. System validates the information
6. System checks for duplicate username and email
7. System creates new faculty account
8. System displays success message
9. Faculty is redirected to login page

Alternative Flows:
6a. If username/email already exists, system displays error message
6b. If validation fails, system displays error message

Postcondition: Faculty account is created and faculty can login'''
    doc.add_paragraph(uc009_content)
    
    # UC-010: Faculty Login
    doc.add_heading('UC-010: Faculty Login', level=2)
    uc010_content = '''Actor: Faculty
Precondition: Faculty has a registered account
Main Flow:
1. Faculty navigates to faculty login page
2. Faculty enters username and password
3. System validates credentials
4. System creates faculty session
5. System redirects to faculty dashboard
6. Faculty can access all faculty features

Alternative Flows:
3a. If credentials are invalid, system displays error message
3b. If account is locked, system displays appropriate message

Postcondition: Faculty is logged in and can access the system'''
    doc.add_paragraph(uc010_content)
    
    # UC-011: View Pending Requests
    doc.add_heading('UC-011: View Pending Requests', level=2)
    uc011_content = '''Actor: Faculty
Precondition: Faculty is logged in
Main Flow:
1. Faculty accesses faculty dashboard
2. System retrieves all booking requests with "Requested" status
3. System displays list of pending requests
4. Faculty can see request details:
   - Student information
   - Room details
   - Event information
   - Booking date and time
5. Faculty can click on a request to view full details

Alternative Flows:
2a. If no pending requests, system displays appropriate message
2b. If database error occurs, system displays error message

Postcondition: Faculty can see all pending booking requests'''
    doc.add_paragraph(uc011_content)
    
    # UC-012: Approve/Reject Booking Request
    doc.add_heading('UC-012: Approve/Reject Booking Request', level=2)
    uc012_content = '''Actor: Faculty
Precondition: Faculty is logged in and viewing booking request
Main Flow:
1. Faculty views detailed booking request
2. Faculty reviews all request information
3. Faculty makes decision (Approve/Reject)
4. Faculty provides reason for decision (optional)
5. Faculty submits decision
6. System validates faculty decision
7. System updates booking status in database
8. System sends notification to student
9. System updates room status if approved
10. System displays confirmation message

Alternative Flows:
6a. If system error occurs, decision is not processed
7a. If room becomes unavailable, system displays conflict message

Postcondition: Booking request status is updated and student is notified'''
    doc.add_paragraph(uc012_content)
    
    # UC-013: View Room Details
    doc.add_heading('UC-013: View Room Details', level=2)
    uc013_content = '''Actor: Faculty
Precondition: Faculty is logged in
Main Flow:
1. Faculty selects a room from the dashboard
2. Faculty clicks on "Show Details" button
3. System retrieves room information and booking details
4. System displays comprehensive room information:
   - Room specifications
   - Current booking status
   - Associated booking requests
   - Room availability schedule
5. Faculty can view all related booking information

Alternative Flows:
3a. If room not found, system displays error message
3b. If database error occurs, system displays error message

Postcondition: Faculty has complete room information for decision making'''
    doc.add_paragraph(uc013_content)
    
    # UC-014: Update Faculty Profile
    doc.add_heading('UC-014: Update Faculty Profile', level=2)
    uc014_content = '''Actor: Faculty
Precondition: Faculty is logged in
Main Flow:
1. Faculty navigates to "Change Information" page
2. System retrieves current faculty information
3. System displays editable form with current data
4. Faculty modifies required fields
5. Faculty submits updated information
6. System validates the updated data
7. System updates faculty profile in database
8. System displays success message

Alternative Flows:
6a. If validation fails, system displays error message
7a. If update fails, system displays error message

Postcondition: Faculty profile is updated with new information'''
    doc.add_paragraph(uc014_content)
    
    # UC-015: Change Faculty Password
    doc.add_heading('UC-015: Change Faculty Password', level=2)
    uc015_content = '''Actor: Faculty
Precondition: Faculty is logged in
Main Flow:
1. Faculty navigates to "Change Password" page
2. Faculty enters current password
3. Faculty enters new password
4. Faculty confirms new password
5. System validates current password
6. System validates new password strength
7. System updates password in database
8. System displays success message
9. Faculty is logged out and must login with new password

Alternative Flows:
5a. If current password is incorrect, system displays error
6a. If new password is weak, system prompts for stronger password
6b. If passwords don't match, system displays error

Postcondition: Faculty password is updated and session is terminated'''
    doc.add_paragraph(uc015_content)
    
    # System Use Cases
    doc.add_heading('System Use Cases', level=1)
    
    # UC-016: Session Management
    doc.add_heading('UC-016: Session Management', level=2)
    uc016_content = '''Actor: System
Precondition: User is logged in
Main Flow:
1. System creates user session upon login
2. System stores session data in MongoDB
3. System tracks session activity
4. System validates session for each request
5. System extends session on activity
6. System expires session after 15 minutes of inactivity
7. System cleans up expired sessions

Alternative Flows:
4a. If session is invalid, system redirects to login page
6a. If user is active, session is extended automatically

Postcondition: Secure session management is maintained'''
    doc.add_paragraph(uc016_content)
    
    # UC-017: Data Validation
    doc.add_heading('UC-017: Data Validation', level=2)
    uc017_content = '''Actor: System
Precondition: User submits form data
Main Flow:
1. System receives form data
2. System validates required fields
3. System validates data format and type
4. System checks for data constraints
5. System validates business rules
6. System returns validation results
7. If valid, system processes data
8. If invalid, system returns error messages

Alternative Flows:
7a. If processing fails, system rolls back transaction
8a. If validation fails, system provides specific error messages

Postcondition: Data integrity is maintained'''
    doc.add_paragraph(uc017_content)
    
    # UC-018: Room Status Management
    doc.add_heading('UC-018: Room Status Management', level=2)
    uc018_content = '''Actor: System
Precondition: Booking request is processed
Main Flow:
1. System receives booking approval/rejection
2. System updates room status in database
3. System checks for conflicts with other bookings
4. System updates booking request status
5. System sends notifications to relevant users
6. System maintains audit trail

Alternative Flows:
3a. If conflict detected, system prevents status change
4a. If update fails, system maintains previous status

Postcondition: Room status is accurately maintained'''
    doc.add_paragraph(uc018_content)
    
    # UC-019: Password Reset
    doc.add_heading('UC-019: Password Reset', level=2)
    uc019_content = '''Actor: Student/Faculty
Precondition: User has forgotten password
Main Flow:
1. User navigates to forgot password page
2. User enters registered email address
3. System validates email address
4. System generates reset token
5. System stores reset token with expiration
6. System sends reset link to email
7. User clicks reset link
8. User enters new password
9. System validates reset token
10. System updates password
11. System invalidates reset token
12. System displays success message

Alternative Flows:
3a. If email not found, system displays error message
9a. If token expired, system displays error message
10a. If password update fails, system displays error message

Postcondition: User password is reset and user can login'''
    doc.add_paragraph(uc019_content)
    
    # UC-020: Contact Form Submission
    doc.add_heading('UC-020: Contact Form Submission', level=2)
    uc020_content = '''Actor: Student
Precondition: Student is logged in
Main Flow:
1. Student navigates to contact page
2. Student fills out contact form with:
   - Subject
   - Message
   - Contact preference
3. Student submits contact form
4. System validates form data
5. System stores contact information
6. System sends notification to administrators
7. System displays confirmation message

Alternative Flows:
4a. If required fields missing, system displays validation error
5a. If storage fails, system displays error message

Postcondition: Contact message is submitted and stored'''
    doc.add_paragraph(uc020_content)
    
    # Use Case Diagram Description
    doc.add_heading('Use Case Relationships', level=1)
    relationships_text = '''The use cases are interconnected through various relationships:

Include Relationships:
• All authenticated use cases include session validation
• All form submissions include data validation
• All profile updates include password verification

Extend Relationships:
• Password reset extends login functionality
• Contact form extends user support functionality
• Feedback extends user experience functionality

Generalization:
• Student and Faculty registration share common validation patterns
• Student and Faculty login share common authentication patterns
• Student and Faculty profile management share common update patterns'''
    doc.add_paragraph(relationships_text)
    
    # Business Rules
    doc.add_heading('Business Rules', level=1)
    business_rules = '''• Students can only book rooms during business hours (8 AM to 8 PM)
• Booking requests must be submitted at least 24 hours in advance
• Students can have maximum 3 pending requests at any time
• Faculty must respond to requests within 48 hours
• Room bookings are limited to 4 hours per session
• Students can only book rooms for their registered society/club events
• Faculty can approve/reject requests based on room availability
• All booking activities are logged for audit purposes
• Session timeout is set to 15 minutes of inactivity
• Password must meet minimum security requirements'''
    doc.add_paragraph(business_rules)
    
    # Non-Functional Requirements
    doc.add_heading('Non-Functional Requirements', level=1)
    nfr_text = '''Performance Requirements:
• System should support up to 200 concurrent users
• Page load time should be under 3 seconds
• Database queries should complete within 1 second

Security Requirements:
• All passwords must be encrypted
• Session data must be securely stored
• User input must be validated and sanitized
• Access control must be enforced for all protected routes

Reliability Requirements:
• System availability should be 99.5%
• Data backup should be performed daily
• System should recover gracefully from failures

Usability Requirements:
• Interface should be intuitive and user-friendly
• System should be accessible on mobile devices
• Error messages should be clear and helpful'''
    doc.add_paragraph(nfr_text)
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion_text = '''This use case documentation provides a comprehensive overview of all interactions between users and the TARMS system. The use cases cover the complete booking workflow from student registration to faculty approval, ensuring all stakeholders understand their roles and responsibilities. The documentation serves as a foundation for system development, testing, and user training.'''
    doc.add_paragraph(conclusion_text)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run('Document Version: 1.0\nLast Updated: 2025\nTARMS Project Documentation')
    footer_run.font.size = Inches(0.1)
    
    # Save the document
    doc.save('/Users/adityasinha/Downloads/TARMS-main/TARMS_UseCase_Documentation.docx')
    print("TARMS Use Case Word document created successfully!")

if __name__ == "__main__":
    create_tarms_usecase_document()


